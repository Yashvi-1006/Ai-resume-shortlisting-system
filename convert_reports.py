"""
Convert Markdown Reports to PDF and DOCX formats
This script reads the markdown technical report and business strategy documents
and converts them to professional PDF and DOCX formats for sharing and printing.
"""

import os
from pathlib import Path

# Try to import required libraries
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    print("✓ python-docx installed")
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    print("✓ reportlab installed")
except ImportError:
    print("Installing reportlab...")
    os.system("pip install reportlab")
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors

def read_markdown_file(filepath):
    """Read markdown file and return content"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def markdown_to_docx(md_content, output_path):
    """Convert markdown content to DOCX format"""
    doc = Document()
    
    lines = md_content.split('\n')
    
    for line in lines:
        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Handle headings
        if line.startswith('# '):
            heading = line.replace('# ', '').strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            heading = line.replace('## ', '').strip()
            p = doc.add_heading(heading, level=2)
        elif line.startswith('### '):
            heading = line.replace('### ', '').strip()
            p = doc.add_heading(heading, level=3)
        elif line.startswith('#### '):
            heading = line.replace('#### ', '').strip()
            p = doc.add_heading(heading, level=4)
        
        # Handle bold and italic
        elif line.startswith('**') or line.startswith('- '):
            # Convert markdown bold/italic to text
            line = line.replace('**', '')
            if line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.bold = True
        
        # Handle code blocks
        elif line.startswith('```'):
            continue  # Skip markdown code markers
        
        # Regular paragraph
        else:
            # Clean up markdown symbols
            clean_line = (line
                         .replace('**', '')
                         .replace('`', '')
                         .replace('|', ' ')
                         .strip())
            
            if clean_line and not clean_line.startswith('#'):
                doc.add_paragraph(clean_line)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'
    
    doc.save(output_path)
    print(f"✓ Created: {output_path}")

def markdown_to_pdf(md_content, output_path):
    """Convert markdown content to PDF format"""
    pdf_doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#0066cc'),
        spaceAfter=12,
        alignment=1  # Center alignment
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#0066cc'),
        spaceAfter=6,
        spaceBefore=6
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )
    
    story = []
    lines = md_content.split('\n')
    
    for line in lines:
        if not line.strip():
            story.append(Spacer(1, 0.1*inch))
            continue
        
        # Handle main title
        if line.startswith('# ') and not line.startswith('## '):
            heading = line.replace('# ', '').strip()
            story.append(Paragraph(heading, title_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Handle secondary headings
        elif line.startswith('## '):
            heading = line.replace('## ', '').strip()
            story.append(Paragraph(heading, heading_style))
        
        # Handle other headings
        elif line.startswith('### ') or line.startswith('#### '):
            heading = line.replace('### ', '').replace('#### ', '').strip()
            # Escape properly for XML
            heading = heading.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(heading, normal_style))
        
        # Handle regular content
        elif line.strip():
            # Properly clean and escape HTML
            clean_line = line.strip()
            # First escape XML special chars except our intentional tags
            clean_line = clean_line.replace('&', '&amp;')
            clean_line = clean_line.replace('<', '&lt;').replace('>', '&gt;')
            # Now handle markdown bold
            clean_line = clean_line.replace('**', '')
            clean_line = clean_line.replace('`', '')
            
            if clean_line and not clean_line.startswith('|'):
                try:
                    story.append(Paragraph(clean_line, normal_style))
                except Exception:
                    # If there's a parsing error, just add plain text
                    story.append(Paragraph(clean_line.replace('&amp;', '&'), normal_style))
    
    try:
        pdf_doc.build(story)
        print(f"✓ Created: {output_path}")
    except Exception as e:
        print(f"✗ Error creating PDF: {e}")
        print("  Note: PDF creation may fail with complex markdown. DOCX format recommended.")

def main():
    """Main conversion function"""
    
    project_dir = Path(r'c:\Users\YASHVI SHAH\OneDrive\Desktop\AI_RESUME_SHORTLISTING_SYSTEM')
    
    # File paths
    tech_report_md = project_dir / 'PROJECT_TECHNICAL_REPORT.md'
    business_report_md = project_dir / 'BUSINESS_VISION_STRATEGY.md'
    
    # Output paths
    tech_report_docx = project_dir / 'PROJECT_TECHNICAL_REPORT.docx'
    tech_report_pdf = project_dir / 'PROJECT_TECHNICAL_REPORT.pdf'
    business_report_docx = project_dir / 'BUSINESS_VISION_STRATEGY.docx'
    business_report_pdf = project_dir / 'BUSINESS_VISION_STRATEGY.pdf'
    
    print("\n" + "="*60)
    print("📄 MARKDOWN TO PDF/DOCX CONVERSION")
    print("="*60 + "\n")
    
    # Convert Technical Report
    if tech_report_md.exists():
        print(f"📖 Processing: {tech_report_md.name}")
        content = read_markdown_file(tech_report_md)
        
        print("\n  Converting to DOCX...")
        markdown_to_docx(content, str(tech_report_docx))
        
        print("  Converting to PDF...")
        markdown_to_pdf(content, str(tech_report_pdf))
    else:
        print(f"✗ File not found: {tech_report_md}")
    
    print()
    
    # Convert Business Report
    if business_report_md.exists():
        print(f"📖 Processing: {business_report_md.name}")
        content = read_markdown_file(business_report_md)
        
        print("\n  Converting to DOCX...")
        markdown_to_docx(content, str(business_report_docx))
        
        print("  Converting to PDF...")
        markdown_to_pdf(content, str(business_report_pdf))
    else:
        print(f"✗ File not found: {business_report_md}")
    
    print("\n" + "="*60)
    print("✅ CONVERSION COMPLETE")
    print("="*60)
    print("\n📂 Output Files:")
    print(f"  • {tech_report_docx.name}")
    print(f"  • {tech_report_pdf.name}")
    print(f"  • {business_report_docx.name}")
    print(f"  • {business_report_pdf.name}")
    print("\n💡 Tip: Use DOCX format for better formatting and editing")
    print("         Use PDF format for sharing and printing\n")

if __name__ == '__main__':
    main()
